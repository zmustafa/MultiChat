from __future__ import annotations

import csv
import io
import os

from .config import settings

# Per-document inline injection budget (characters) when stuffing extracted text
# directly into the model prompt. Larger docs are truncated; the model can call the
# read_document tool to retrieve the full text.
INLINE_BUDGET = 8000


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001
            pages.append("")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out: list[str] = []
    for ws in wb.worksheets:
        out.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                out.append("\t".join(cells))
    return "\n".join(out)


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    return "\n".join("\t".join(row) for row in reader)


def _extract_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


# mime/extension -> extractor
_EXTRACTORS = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": _extract_xlsx,
    "text/csv": _extract_csv,
    "text/plain": _extract_text,
    "text/markdown": _extract_text,
}

_EXT_MIME = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".csv": "text/csv",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


def is_document(mime_type: str, filename: str) -> bool:
    if mime_type in _EXTRACTORS:
        return True
    ext = os.path.splitext(filename or "")[1].lower()
    return ext in _EXT_MIME


def normalize_doc_mime(mime_type: str, filename: str) -> str:
    if mime_type in _EXTRACTORS:
        return mime_type
    ext = os.path.splitext(filename or "")[1].lower()
    return _EXT_MIME.get(ext, mime_type or "application/octet-stream")


def extract_text(data: bytes, mime_type: str, filename: str) -> str:
    """Extract plain text from a document's bytes. Returns "" if unsupported/failed."""
    mime = normalize_doc_mime(mime_type, filename)
    extractor = _EXTRACTORS.get(mime)
    if not extractor:
        return ""
    try:
        return (extractor(data) or "").strip()
    except Exception:  # noqa: BLE001
        return ""


def document_prompt_block(attachments) -> str:
    """Build an inline text block of the extracted content of document attachments,
    to append to a user message so the model can read them directly."""
    blocks: list[str] = []
    for att in attachments:
        if att.kind != "document" or not att.extracted_text:
            continue
        text = att.extracted_text
        truncated = len(text) > INLINE_BUDGET
        snippet = text[:INLINE_BUDGET]
        note = (
            f"\n\n[…truncated. Call read_document(name=\"{att.filename}\") "
            "for the full text.]"
            if truncated
            else ""
        )
        blocks.append(
            f'--- Attached document: "{att.filename}" ---\n{snippet}{note}\n'
            f'--- end of "{att.filename}" ---'
        )
    if not blocks:
        return ""
    return "\n\n" + "\n\n".join(blocks)
