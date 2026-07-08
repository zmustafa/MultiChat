from __future__ import annotations

import io
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from .artifacts import (
    download_result,
    generated_dir,
    new_stored_name,
    resolve_image_bytes,
    safe_download_name,
)
from .base import ToolContext, ToolDef, ToolResult


class DocxGenerateTool:
    definition = ToolDef(
        name="generate_docx",
        description=(
            "Create a downloadable Microsoft Word (.docx) document from a structured "
            "outline that you author. Use this whenever the user asks for a Word "
            "document, report, letter, or write-up. Provide a title and a list of "
            "sections (each with an optional heading, paragraphs, bullet points, and/or "
            "a table). Returns a Markdown download link — include it verbatim in your "
            "reply."
        ),
        parameters={
            "type": "object",
            "properties": {
                "title": {"type": "string", "description": "Document title."},
                "subtitle": {"type": "string", "description": "Optional subtitle/byline."},
                "sections": {
                    "type": "array",
                    "description": "The document sections, in order.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "heading": {
                                "type": "string",
                                "description": "Optional section heading.",
                            },
                            "paragraphs": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Body paragraphs.",
                            },
                            "bullets": {
                                "type": "array",
                                "items": {"type": "string"},
                                "description": "Bullet points.",
                            },
                            "table": {
                                "type": "object",
                                "description": "Optional table.",
                                "properties": {
                                    "columns": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                    },
                                    "rows": {
                                        "type": "array",
                                        "items": {
                                            "type": "array",
                                            "items": {"type": "string"},
                                        },
                                    },
                                },
                            },
                            "image": {
                                "type": "string",
                                "description": (
                                    "Optional image to embed: a data: URI, an http(s) "
                                    "image URL, or a generated /api/files/<name> link."
                                ),
                            },
                            "caption": {
                                "type": "string",
                                "description": "Optional caption shown under the image.",
                            },
                        },
                    },
                },
            },
            "required": ["title", "sections"],
        },
    )

    async def run(self, args: dict, ctx: ToolContext) -> ToolResult:
        title = (args.get("title") or "").strip() or "Document"
        subtitle = (args.get("subtitle") or "").strip()
        sections = args.get("sections") or []
        if not isinstance(sections, list) or not sections:
            return ToolResult(content="No content was provided.", citations=None)

        try:
            doc = Document()
            h = doc.add_heading(title, level=0)
            h.runs[0].font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
            if subtitle:
                p = doc.add_paragraph()
                run = p.add_run(subtitle)
                run.italic = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            for sec in sections:
                if not isinstance(sec, dict):
                    continue
                heading = (sec.get("heading") or "").strip()
                if heading:
                    doc.add_heading(heading, level=1)
                paragraphs = sec.get("paragraphs") or []
                if isinstance(paragraphs, str):
                    paragraphs = [paragraphs]
                for para in paragraphs:
                    text = str(para).strip()
                    if text:
                        doc.add_paragraph(text)
                bullets = sec.get("bullets") or []
                if isinstance(bullets, str):
                    bullets = [bullets]
                for b in bullets:
                    text = str(b).strip()
                    if text:
                        doc.add_paragraph(text, style="List Bullet")

                table = sec.get("table")
                if isinstance(table, dict):
                    cols = table.get("columns") or []
                    rows = table.get("rows") or []
                    if cols:
                        t = doc.add_table(rows=1, cols=len(cols))
                        t.style = "Light Grid Accent 1"
                        for i, c in enumerate(cols):
                            t.rows[0].cells[i].text = str(c)
                        for r in rows:
                            cells = t.add_row().cells
                            for i in range(len(cols)):
                                cells[i].text = str(r[i]) if i < len(r) else ""

                image_ref = (sec.get("image") or "").strip()
                if image_ref:
                    data = resolve_image_bytes(image_ref)
                    if data:
                        try:
                            doc.add_picture(io.BytesIO(data), width=Inches(5.5))
                        except Exception:  # noqa: BLE001
                            pass
                        else:
                            caption = (sec.get("caption") or "").strip()
                            if caption:
                                cap = doc.add_paragraph()
                                run = cap.add_run(caption)
                                run.italic = True
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

            stored_name = new_stored_name("docx")
            path = os.path.join(generated_dir(), stored_name)
            doc.save(path)
        except Exception as exc:  # noqa: BLE001
            return ToolResult(content=f"Failed to build document: {exc}", citations=None)

        download_name = safe_download_name(title, "docx")
        return download_result(
            stored_name, download_name, f"Created a Word document with {len(sections)} section(s)."
        )
