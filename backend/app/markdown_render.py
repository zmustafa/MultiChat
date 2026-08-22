"""Render Markdown text into real Word (python-docx) and PDF (reportlab) elements.

LLM answers are Markdown. Exporting them by dumping the raw string leaves literal
``**bold**``, ``# heading``, ``- list`` and fenced code markers in the Word/PDF output.
This module parses the common Markdown constructs LLMs actually emit — headings, bold/
italic/inline-code, links, bullet/numbered lists, fenced code blocks, block quotes,
horizontal rules and pipe tables — into a small list of block tokens, then renders those
tokens with proper document formatting for each backend.
"""
from __future__ import annotations

import io
import itertools
import os
import re
import unicodedata
from collections.abc import Callable, Iterable, Sequence
from typing import Any
from xml.sax.saxutils import escape

# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------

# A block token is a tuple whose first element names its kind:
#   ("h", level:int, text:str)
#   ("p", text:str)
#   ("ul", items:list[str])       ("ol", items:list[str], start:int)
#   ("code", text:str, lang:str)
#   ("quote", blocks:list[Block])
#   ("hr",)
#   ("table", cols:list[str], rows:list[list[str]])
#   ("image", alt:str, src:str)
Block = tuple

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_HR_RE = re.compile(r"^(\*\s*){3,}$|^(-\s*){3,}$|^(_\s*){3,}$")
_UL_RE = re.compile(r"^\s*[-*+]\s+")
_OL_RE = re.compile(r"^\s*\d+[.)]\s+")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
_IMAGE_ONLY_RE = re.compile(r'^!\[([^\]]*)\]\(\s*(\S+?)(?:\s+"[^"]*")?\s*\)$')


def _is_block_start(line: str) -> bool:
    s = line.strip()
    if not s:
        return True
    if s.startswith("```") or s.startswith("~~~"):
        return True
    if _HEADING_RE.match(s):
        return True
    if _UL_RE.match(line) or _OL_RE.match(line):
        return True
    if s.startswith(">"):
        return True
    if _HR_RE.match(s):
        return True
    return False


def _absorb_continuation(
    lines: list[str], i: int, n: int, text: str
) -> tuple[int, str]:
    """Fold a list item's soft-wrapped continuation lines into its text.

    A bullet whose text wraps onto the next (unmarked) source line is ONE item; without
    this the wrapped line became a separate paragraph rendered back at the left margin.
    """
    while i < n and lines[i].strip() and not _is_block_start(lines[i]):
        text = f"{text} {lines[i].strip()}".strip()
        i += 1
    return i, text


def _split_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def parse_blocks(md: str) -> list[Block]:
    lines = (md or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[Block] = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```") or stripped.startswith("~~~"):
            fence = stripped[:3]
            lang = stripped[3:].strip().split()[0].lower() if stripped[3:].strip() else ""
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].strip().startswith(fence):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            blocks.append(("code", "\n".join(code_lines), lang))
            continue

        if not stripped:
            i += 1
            continue

        m = _HEADING_RE.match(stripped)
        if m:
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            i += 1
            continue

        if _HR_RE.match(stripped):
            blocks.append(("hr",))
            i += 1
            continue

        # A line that is nothing but an image, e.g. a chart produced by a tool.
        m_img = _IMAGE_ONLY_RE.match(stripped)
        if m_img:
            blocks.append(("image", m_img.group(1), m_img.group(2)))
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < n:
                cur = lines[i]
                if cur.strip().startswith(">"):
                    quote_lines.append(re.sub(r"^\s*>\s?", "", cur))
                    i += 1
                    continue
                # Lazy continuation: an unmarked line that is not itself a new block still
                # belongs to the paragraph the quote was in the middle of.
                if cur.strip() and not _is_block_start(cur):
                    quote_lines.append(cur.strip())
                    i += 1
                    continue
                break
            # A quote is a container, not a string: its headings, lists, tables and code
            # fences are real blocks and must render as such.
            blocks.append(("quote", parse_blocks("\n".join(quote_lines))))
            continue

        # Pipe table: a header row followed by a |---|---| separator row
        if "|" in line and i + 1 < n and _TABLE_SEP_RE.match(lines[i + 1]):
            header = _split_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < n and lines[i].strip() and "|" in lines[i]:
                rows.append(_split_row(lines[i]))
                i += 1
            blocks.append(("table", header, rows))
            continue

        if _UL_RE.match(line):
            items: list[str] = []
            while i < n and _UL_RE.match(lines[i]):
                text = re.sub(r"^\s*[-*+]\s+", "", lines[i]).strip()
                i += 1
                i, text = _absorb_continuation(lines, i, n, text)
                items.append(text)
            blocks.append(("ul", items))
            continue

        if _OL_RE.match(line):
            # Keep the author's numbering: a list interrupted by a paragraph continues
            # counting instead of restarting at 1.
            first = re.match(r"^\s*(\d+)", line)
            start = int(first.group(1)) if first else 1
            items = []
            while i < n and _OL_RE.match(lines[i]):
                text = re.sub(r"^\s*\d+[.)]\s+", "", lines[i]).strip()
                i += 1
                i, text = _absorb_continuation(lines, i, n, text)
                items.append(text)
            blocks.append(("ol", items, start))
            continue

        # Paragraph: gather soft-wrapped lines until a blank line or a new block.
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not _is_block_start(lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append(("p", " ".join(para_lines)))

    return blocks


# ---------------------------------------------------------------------------
# Inline parsing (bold / italic / code / links)
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"`([^`]+)`"                      # 1 code
    r"|\*\*([^*]+)\*\*"              # 2 bold
    r"|__([^_]+)__"                  # 3 bold
    r"|\*([^*]+)\*"                  # 4 italic
    r"|(?<![A-Za-z0-9])_([^_]+)_(?![A-Za-z0-9])"  # 5 italic
    r"|\[([^\]]+)\]\(([^)\s]+)\)"    # 6 link text, 7 href
)


def _inline_spans(text: str) -> list[tuple[str, set[str], str | None]]:
    """Split inline Markdown into (text, styles, href) spans."""
    spans: list[tuple[str, set[str], str | None]] = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            spans.append((text[pos:m.start()], set(), None))
        if m.group(1) is not None:
            spans.append((m.group(1), {"code"}, None))
        elif m.group(2) is not None:
            spans.append((m.group(2), {"bold"}, None))
        elif m.group(3) is not None:
            spans.append((m.group(3), {"bold"}, None))
        elif m.group(4) is not None:
            spans.append((m.group(4), {"italic"}, None))
        elif m.group(5) is not None:
            spans.append((m.group(5), {"italic"}, None))
        elif m.group(6) is not None:
            spans.append((m.group(6), {"link"}, m.group(7)))
        pos = m.end()
    if pos < len(text):
        spans.append((text[pos:], set(), None))
    return spans or [(text, set(), None)]


def _strip_inline(text: str) -> str:
    return "".join(s[0] for s in _inline_spans(text))


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------


def _add_inline_docx(paragraph: Any, text: str) -> None:
    from docx.shared import Pt, RGBColor

    for content, styles, href in _inline_spans(text):
        if not content:
            continue
        run = paragraph.add_run(content)
        if "bold" in styles:
            run.bold = True
        if "italic" in styles:
            run.italic = True
        if "code" in styles:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        if href or "link" in styles:
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            run.underline = True


def render_markdown_docx(
    doc: Any,
    md: str,
    base_level: int = 2,
    diagrams: Sequence[dict] | None = None,
) -> None:
    """Append Markdown ``md`` to a python-docx ``doc`` as formatted elements.

    ``base_level`` is the docx heading level that a Markdown ``#`` maps to, so content
    headings nest below the surrounding section headings.

    ``diagrams`` optionally supplies pre-rendered PNGs for ```mermaid``` fences (the chat UI
    rasterizes what it is already showing). Without one, a fence falls back to its source
    text — readable, but not a diagram.
    """
    import io

    from docx.shared import Emu, Inches

    pending = [dict(d) for d in (diagrams or [])]

    def take_diagram(code: str) -> dict | None:
        key = (code or "").strip()
        for i, d in enumerate(pending):
            if (d.get("code") or "").strip() == key:
                return pending.pop(i)
        return None

    def place_diagram(data: bytes) -> bool:
        try:
            shape = doc.add_picture(io.BytesIO(data))
        except Exception:  # noqa: BLE001 — fall back to the source text
            return False
        max_w, max_h = Inches(6.0), Inches(4.5)
        if shape.width > max_w:
            shape.height = Emu(int(shape.height * max_w / shape.width))
            shape.width = max_w
        if shape.height > max_h:
            shape.width = Emu(int(shape.width * max_h / shape.height))
            shape.height = max_h
        return True

    _render_blocks_docx(
        doc, parse_blocks(md), base_level, take_diagram, place_diagram
    )


def _render_blocks_docx(
    doc: Any,
    blocks: Sequence[Block],
    base_level: int,
    take_diagram: Callable[[str], dict | None],
    place_diagram: Callable[[bytes], bool],
    indent: float = 0.0,
) -> None:
    from docx.shared import Inches, Pt, RGBColor

    def para(style: str | None = None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        return p

    for block in blocks:
        kind = block[0]
        if kind == "h":
            level = min(base_level + block[1] - 1, 9)
            hp = doc.add_heading("", level=level)
            if indent:
                hp.paragraph_format.left_indent = Inches(indent)
            _add_inline_docx(hp, block[2])
        elif kind == "p":
            _add_inline_docx(para(), block[1])
        elif kind == "ul":
            for it in block[1]:
                _add_inline_docx(para("List Bullet"), it)
        elif kind == "ol":
            for it in block[1]:
                _add_inline_docx(para("List Number"), it)
        elif kind == "code":
            lang = block[2] if len(block) > 2 else ""
            if lang == "mermaid":
                d = take_diagram(block[1])
                if d and d.get("data") and place_diagram(d["data"]):
                    continue
            for cl in block[1].split("\n"):
                p = para()
                run = p.add_run(cl or " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif kind == "quote":
            # Quoted content keeps its own block structure; the quote only adds an indent.
            _render_blocks_docx(
                doc, block[1], base_level, take_diagram, place_diagram, indent + 0.3
            )
        elif kind == "hr":
            para().add_run("─" * 30)
        elif kind == "table":
            cols, rows = block[1], block[2]
            if not cols:
                continue
            t = doc.add_table(rows=1, cols=len(cols))
            try:
                t.style = "Light Grid Accent 1"
            except Exception:  # noqa: BLE001
                pass
            for c_i, c in enumerate(cols):
                t.rows[0].cells[c_i].text = _strip_inline(str(c))
            for r in rows:
                cells = t.add_row().cells
                for c_i in range(len(cols)):
                    cells[c_i].text = _strip_inline(str(r[c_i])) if c_i < len(r) else ""


# ---------------------------------------------------------------------------
# PDF rendering (reportlab)
# ---------------------------------------------------------------------------

# GitHub-dark palette — the same theme the chat lanes use for fenced code blocks.
CODE_BG = "#0D1117"
CODE_FG = "#E6EDF3"

_FONTS: dict[str, str] | None = None
_GLYPH_OK: Callable[[str], bool] | None = None

# Characters LLM answers use constantly that no PDF base font can draw. Anything not
# listed and not drawable is dropped rather than rendered as a "missing glyph" box.
_CHAR_FALLBACKS = {
    "\u2192": "->", "\u2190": "<-", "\u2191": "^", "\u2193": "v", "\u2194": "<->",
    "\u21d2": "=>", "\u21d0": "<=", "\u21d4": "<=>",
    "\u2713": "[x]", "\u2714": "[x]", "\u2705": "[x]", "\u2611": "[x]",
    "\u2717": "[ ]", "\u2718": "[ ]", "\u274c": "[!]", "\u2b55": "( )", "\u2b1c": "[ ]",
    "\u26a0": "!", "\u2139": "i", "\u2757": "!", "\u2753": "?",
    "\u2605": "*", "\u2606": "*", "\u2b50": "*", "\u25cf": "\u2022", "\u25cb": "\u2022",
    "\u25aa": "\u2022", "\u25ab": "\u2022", "\u25e6": "\u2022", "\u2023": "\u2022",
    "\u25b6": ">", "\u25c0": "<", "\u25b2": "^", "\u25bc": "v",
    "\u2260": "!=", "\u2248": "~", "\u221e": "inf", "\u2261": "==",
    "\u2265": ">=", "\u2264": "<=", "\u2212": "-", "\u2044": "/",
    "\u2500": "-", "\u2501": "-", "\u2502": "|", "\u2503": "|", "\u2550": "=", "\u2551": "|",
    "\u250c": "+", "\u2510": "+", "\u2514": "+", "\u2518": "+", "\u251c": "+",
    "\u2524": "+", "\u252c": "+", "\u2534": "+", "\u253c": "+",
    "\u00a0": " ", "\u202f": " ", "\u2009": " ", "\u200b": "", "\ufe0f": "", "\u2060": "",
}


def _font_dirs() -> list[str]:
    import reportlab

    return [
        "/usr/share/fonts/truetype/dejavu",
        "/usr/share/fonts/dejavu",
        "/usr/share/fonts/TTF",
        "/usr/local/share/fonts",
        "/Library/Fonts",
        "C:/Windows/Fonts",
        os.path.join(os.path.dirname(reportlab.__file__), "fonts"),
    ]


def pdf_fonts() -> dict[str, str]:
    """Font names to use for PDF body/bold/italic/mono text.

    Prefers DejaVu (broad Unicode coverage) when the host provides it, so arrows, box
    drawing and accented text survive; otherwise falls back to the built-in Helvetica and
    Courier, and :func:`pdf_safe` transliterates whatever the font cannot draw.
    """
    global _FONTS
    if _FONTS is not None:
        return _FONTS

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fonts = {
        "body": "Helvetica",
        "bold": "Helvetica-Bold",
        "italic": "Helvetica-Oblique",
        "boldItalic": "Helvetica-BoldOblique",
        "mono": "Courier",
        "monoBold": "Courier-Bold",
    }
    dirs = _font_dirs()

    def find(filename: str) -> str | None:
        for d in dirs:
            path = os.path.join(d, filename)
            if os.path.exists(path):
                return path
        return None

    def register_family(prefix: str, files: dict[str, str], keys: Sequence[str]) -> bool:
        paths = {k: find(files[k]) for k in keys}
        if not all(paths.values()):
            return False
        names = {k: f"{prefix}-{k}" for k in keys}
        for k in keys:
            pdfmetrics.registerFont(TTFont(names[k], paths[k]))
        pdfmetrics.registerFontFamily(
            names["body"],
            normal=names["body"],
            bold=names.get("bold", names["body"]),
            italic=names.get("italic", names["body"]),
            boldItalic=names.get("boldItalic", names["body"]),
        )
        fonts.update(names)
        return True

    try:
        register_family(
            "MCSans",
            {
                "body": "DejaVuSans.ttf",
                "bold": "DejaVuSans-Bold.ttf",
                "italic": "DejaVuSans-Oblique.ttf",
                "boldItalic": "DejaVuSans-BoldOblique.ttf",
            },
            ("body", "bold", "italic", "boldItalic"),
        )
    except Exception:  # noqa: BLE001 — a broken font file must not break exporting
        pass
    try:
        mono, mono_bold = find("DejaVuSansMono.ttf"), find("DejaVuSansMono-Bold.ttf")
        if mono and mono_bold:
            pdfmetrics.registerFont(TTFont("MCMono", mono))
            pdfmetrics.registerFont(TTFont("MCMono-Bold", mono_bold))
            pdfmetrics.registerFontFamily(
                "MCMono", normal="MCMono", bold="MCMono-Bold",
                italic="MCMono", boldItalic="MCMono-Bold",
            )
            fonts["mono"], fonts["monoBold"] = "MCMono", "MCMono-Bold"
    except Exception:  # noqa: BLE001
        pass

    _FONTS = fonts
    return fonts


def _glyph_ok() -> Callable[[str], bool]:
    global _GLYPH_OK
    if _GLYPH_OK is not None:
        return _GLYPH_OK

    from reportlab.pdfbase import pdfmetrics

    try:
        face = pdfmetrics.getFont(pdf_fonts()["body"]).face
    except Exception:  # noqa: BLE001
        face = None
    table = getattr(face, "charToGlyph", None)
    if table is not None:
        def ok(ch: str) -> bool:
            return ord(ch) in table
    else:
        def ok(ch: str) -> bool:
            try:
                ch.encode("cp1252")
                return True
            except UnicodeEncodeError:
                return False
    _GLYPH_OK = ok
    return ok


def pdf_safe(text: str) -> str:
    """Make ``text`` renderable by the active PDF font.

    Characters the font cannot draw are transliterated (``->`` for an arrow, ``[x]`` for a
    check mark, accents stripped) or dropped, so the PDF never shows missing-glyph boxes.
    """
    if not text:
        return text
    ok = _glyph_ok()
    if all(ch in "\n\t" or ok(ch) for ch in text):
        return text
    out: list[str] = []
    for ch in text:
        if ch in "\n\t" or ok(ch):
            out.append(ch)
            continue
        repl = _CHAR_FALLBACKS.get(ch)
        if repl is None:
            # "ā" -> "a": keep the base letter of a decomposable character.
            repl = "".join(c for c in unicodedata.normalize("NFKD", ch) if not unicodedata.combining(c))
            if repl == ch:
                repl = ""
        out.append("".join(c for c in repl if ok(c)))
    return "".join(out)


def footer_canvas(
    footer_left: str,
    font: str,
    pagesize: tuple[float, float] | None = None,
    margin: float = 54.0,
    attribution: bool = True,
):
    """Canvas subclass stamping a rule, context text, repo link and "Page X of Y".

    The page count is only known once the whole story is laid out, so pages are buffered
    and replayed on save. The centre slot carries a clickable link to the project repo
    (``settings.APP_REPO_URL``); pass ``attribution=False`` to leave it out.
    """
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfgen import canvas as pdf_canvas

    from .config import settings

    page_w, _page_h = pagesize or LETTER
    repo_url = (settings.APP_REPO_URL or "").strip() if attribution else ""
    # Show the bare host/path; the full URL stays behind the link.
    repo_label = pdf_safe(re.sub(r"^https?://", "", repo_url)) if repo_url else ""
    # Pages are replayed from buffered state on save, which rewinds reportlab's annotation
    # counter - every footer link would then be named "Annot.NUMBER1" and the second one
    # would raise "redefining named object". This counter keeps the names unique; the high
    # base keeps it clear of the names used by links inside the document body.
    link_seq = itertools.count(1_000_000)

    class FooterCanvas(pdf_canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._pages: list[dict] = []

        def showPage(self):  # noqa: N802 — reportlab API
            self._pages.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._pages)
            for state in self._pages:
                self.__dict__.update(state)
                self._stamp(total)
                super().showPage()
            super().save()

        def _stamp(self, total: int) -> None:
            self.saveState()
            self.setStrokeColor(HexColor("#E2E8F0"))
            self.setLineWidth(0.5)
            self.line(margin, 44, page_w - margin, 44)
            self.setFont(font, 7.5)
            self.setFillColor(HexColor("#94A3B8"))
            # The left slot is capped to a third of the usable width so a long title
            # can never collide with the centred link.
            usable = page_w - 2 * margin
            self._draw_clipped(pdf_safe(footer_left), margin, 32, usable / 3.0)
            self.drawRightString(page_w - margin, 32, f"Page {self._pageNumber} of {total}")
            if repo_label:
                self._draw_link(page_w / 2.0, 32)
            self.restoreState()

        def _draw_clipped(self, text: str, x: float, y: float, max_width: float) -> None:
            while text and pdfmetrics.stringWidth(text, font, 7.5) > max_width:
                text = text[:-1]
            self.drawString(x, y, text)

        def _draw_link(self, cx: float, y: float) -> None:
            self.setFillColor(HexColor("#6366F1"))
            self.drawCentredString(cx, y, repo_label)
            width = pdfmetrics.stringWidth(repo_label, font, 7.5)
            self._annotationCount = next(link_seq)
            self.linkURL(repo_url, (cx - width / 2, y - 2, cx + width / 2, y + 8), relative=0)
            self.setFillColor(HexColor("#94A3B8"))

    return FooterCanvas


def _inline_pdf(text: str) -> str:
    """Convert inline Markdown to reportlab's mini-HTML markup (fully escaped)."""
    fonts = pdf_fonts()
    out: list[str] = []
    for content, styles, href in _inline_spans(text):
        seg = escape(pdf_safe(content))
        if "code" in styles:
            seg = f'<font face="{fonts["mono"]}" color="#B91C1C">{seg}</font>'
        if "bold" in styles:
            seg = f"<b>{seg}</b>"
        if "italic" in styles:
            seg = f"<i>{seg}</i>"
        if href:
            safe_href = escape(href, {'"': "&quot;"})
            seg = f'<link href="{safe_href}"><font color="#4F46E5">{seg}</font></link>'
        out.append(seg)
    return "".join(out)


# ---------------------------------------------------------------------------
# Fenced code blocks: syntax highlighting
# ---------------------------------------------------------------------------

_PALETTE: dict | None = None


def _palette() -> dict:
    """GitHub-dark token colours keyed by Pygments token type."""
    global _PALETTE
    if _PALETTE is None:
        from pygments.token import Token

        _PALETTE = {
            Token.Comment: "#8B949E",
            Token.Keyword: "#FF7B72",
            Token.Keyword.Type: "#FFA657",
            Token.Operator: "#FF7B72",
            Token.Operator.Word: "#FF7B72",
            Token.Name: CODE_FG,
            Token.Name.Builtin: "#79C0FF",
            Token.Name.Builtin.Pseudo: "#79C0FF",
            Token.Name.Function: "#D2A8FF",
            Token.Name.Class: "#FFA657",
            Token.Name.Namespace: "#FFA657",
            Token.Name.Decorator: "#D2A8FF",
            Token.Name.Tag: "#7EE787",
            Token.Name.Attribute: "#79C0FF",
            Token.Name.Variable: "#FFA657",
            Token.Name.Constant: "#79C0FF",
            Token.Literal: "#A5D6FF",
            Token.String: "#A5D6FF",
            Token.String.Escape: "#79C0FF",
            Token.Number: "#79C0FF",
            Token.Generic.Deleted: "#FFA198",
            Token.Generic.Inserted: "#7EE787",
            Token.Generic.Heading: "#79C0FF",
            Token.Generic.Emph: "#E6EDF3",
            Token.Generic.Prompt: "#8B949E",
            Token.Error: "#FFA198",
        }
    return _PALETTE


def _code_spans(code: str, lang: str) -> list[list[tuple[str, str | None]]]:
    """Tokenize ``code`` into one list of ``(text, colour)`` spans per line."""
    try:
        from pygments import lex
        from pygments.lexers import get_lexer_by_name, guess_lexer

        palette = _palette()
        try:
            lexer = get_lexer_by_name(lang, stripnl=False, stripall=False)
        except Exception:  # noqa: BLE001 — unknown/absent language
            lexer = guess_lexer(code, stripnl=False) if code.strip() else None
        if lexer is None:
            raise ValueError("no lexer")

        def colour(ttype) -> str | None:
            t = ttype
            while t is not None:
                if t in palette:
                    return palette[t]
                t = t.parent
            return None

        lines: list[list[tuple[str, str | None]]] = [[]]
        for ttype, value in lex(code, lexer):
            col = colour(ttype)
            parts = value.split("\n")
            for idx, part in enumerate(parts):
                if idx:
                    lines.append([])
                if part:
                    lines[-1].append((part, col))
        while lines and not lines[-1]:
            lines.pop()
        return lines or [[]]
    except Exception:  # noqa: BLE001 — highlighting is best-effort
        return [[(line, None)] for line in code.split("\n")]


def _wrap_code_spans(
    lines: list[list[tuple[str, str | None]]], max_chars: int
) -> list[list[tuple[str, str | None]]]:
    """Hard-wrap highlighted lines so long code never runs off the page."""
    max_chars = max(20, max_chars)
    wrapped: list[list[tuple[str, str | None]]] = []
    for spans in lines:
        plain = "".join(t for t, _ in spans)
        if len(plain) <= max_chars:
            wrapped.append(spans)
            continue
        colours: list[str | None] = []
        for text, col in spans:
            colours.extend([col] * len(text))
        indent = len(plain) - len(plain.lstrip(" "))
        prefix = " " * min(indent + 2, max_chars // 2)
        pos, first = 0, True
        while pos < len(plain):
            avail = max(12, max_chars - (0 if first else len(prefix)))
            chunk = plain[pos : pos + avail]
            if pos + avail < len(plain):
                brk = max(chunk.rfind(" "), chunk.rfind(","), chunk.rfind(";"))
                if brk > avail * 0.55:
                    chunk = chunk[: brk + 1]
            row: list[tuple[str, str | None]] = []
            if not first:
                row.append((prefix, None))
            for offset, ch in enumerate(chunk):
                col = colours[pos + offset]
                if row and row[-1][1] == col and (first or offset):
                    row[-1] = (row[-1][0] + ch, col)
                else:
                    row.append((ch, col))
            wrapped.append(row)
            pos += len(chunk)
            first = False
    return wrapped


def _code_markup(lines: Iterable[list[tuple[str, str | None]]]) -> str:
    rows: list[str] = []
    for spans in lines:
        parts = []
        for text, colour in spans:
            seg = escape(pdf_safe(text))
            parts.append(f'<font color="{colour}">{seg}</font>' if colour else seg)
        rows.append("".join(parts) or " ")
    return "\n".join(rows)


# ---------------------------------------------------------------------------
# Diagrams and images
# ---------------------------------------------------------------------------


def _diagram_card(data: bytes, max_width: float, max_height: float, natural_pt: float = 0.0):
    """A bordered image card that shrinks to fit the space left on the page.

    A plain reportlab ``Image`` that does not fit simply jumps to the next page, which for
    a large diagram can leave most of a page blank. This flowable scales down instead —
    but only while the result stays readable, otherwise it moves on as usual.
    """
    from reportlab.lib import colors
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Flowable

    reader = ImageReader(io.BytesIO(data))
    px_w, px_h = reader.getSize()
    if not px_w or not px_h:
        return None
    # Rasterized diagrams arrive at 2-4x, so pixel size alone would be enormous. Prefer the
    # diagram's own layout size (scaled up a little so it reads well in print) when known.
    target = natural_pt * 1.6 if natural_pt else max_width
    width = max(1.0, min(max_width, target))
    height = width * px_h / px_w
    if height > max_height:
        height = max_height
        width = height * px_w / px_h

    class DiagramCard(Flowable):
        pad = 8.0
        # Below this the shrunk diagram stops being legible - take the page break instead.
        min_fit = 200.0

        def __init__(self) -> None:
            super().__init__()
            self.hAlign = "CENTER"
            self.spaceBefore = 6
            self.spaceAfter = 12
            self._w, self._h = width + 2 * self.pad, height + 2 * self.pad
            self._iw, self._ih = width, height

        def wrap(self, avail_width: float, avail_height: float):  # noqa: D102
            w, h = width, height
            room_h = avail_height - 2 * self.pad
            if h > room_h >= self.min_fit:
                w, h = w * room_h / h, room_h
            room_w = avail_width - 2 * self.pad
            if w > room_w:
                w, h = room_w, h * room_w / w
            self._iw, self._ih = w, h
            self._w, self._h = w + 2 * self.pad, h + 2 * self.pad
            return self._w, self._h

        def draw(self) -> None:  # noqa: D102
            c = self.canv
            c.saveState()
            c.setFillColor(colors.white)
            c.setStrokeColor(colors.HexColor("#E2E8F0"))
            c.setLineWidth(0.6)
            c.roundRect(0, 0, self._w, self._h, 4, stroke=1, fill=1)
            c.drawImage(
                ImageReader(io.BytesIO(data)), self.pad, self.pad,
                width=self._iw, height=self._ih, mask="auto",
            )
            c.restoreState()

    return DiagramCard()


def image_flowable(data: bytes, max_width: float, max_height: float):
    """Embed arbitrary image bytes (a prompt attachment, say) as a bordered card.

    Returns None when the bytes aren't a readable image, so callers can fall back to
    naming the file instead of failing the whole export.
    """
    try:
        return _diagram_card(data, max_width, max_height)
    except Exception:  # noqa: BLE001 — an unreadable image must not sink the document
        return None


def _fit_col_widths(
    cols: list[str], rows: list[list[str]], content_width: float, font_size: float = 8.5
) -> list[float]:
    """Distribute the frame width across table columns.

    Every column first reserves enough room for its widest unbreakable word (so short
    values like ``99.95%`` are never split across lines), then the slack is shared out in
    proportion to how much text each column holds.
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth

    fonts = pdf_fonts()
    pad = 12.0
    ceiling = content_width / 2
    mins: list[float] = []
    weights: list[float] = []
    for i in range(len(cols)):
        cells = [str(cols[i])] + [str(r[i]) if i < len(r) else "" for r in rows]
        texts = [_strip_inline(c) for c in cells]
        widest_word = 0.0
        for text in texts:
            for word in text.split() or [""]:
                widest_word = max(
                    widest_word, stringWidth(pdf_safe(word), fonts["bold"], font_size)
                )
        mins.append(min(ceiling, widest_word + pad))
        weights.append(float(max(4, min(max((len(t) for t in texts), default=1), 60))))

    total_min = sum(mins)
    if total_min >= content_width:
        return [w * content_width / total_min for w in mins]
    slack = content_width - total_min
    total_weight = sum(weights) or 1.0
    return [m + slack * w / total_weight for m, w in zip(mins, weights, strict=False)]


def markdown_pdf_flowables(
    md: str,
    body_style: Any,
    *,
    diagrams: Sequence[dict] | None = None,
    content_width: float = 6.9 * 72,
) -> list:
    """Return a list of reportlab flowables rendering Markdown ``md``.

    ``diagrams`` optionally supplies pre-rendered images for ```mermaid``` fences, each
    ``{"code": str, "data": bytes, "width": float, "height": float}`` — the chat UI
    rasterizes the diagram it is already showing so the PDF matches the lane exactly.
    """
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        KeepTogether,
        ListFlowable,
        ListItem,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        XPreformatted,
    )
    from reportlab.platypus.flowables import HRFlowable

    from .tools.artifacts import resolve_image_bytes

    fonts = pdf_fonts()
    heading_style = ParagraphStyle(
        "MdHeading", parent=body_style, fontName=fonts["bold"],
        textColor=colors.HexColor("#1E1B4B"), spaceBefore=10, spaceAfter=4,
        keepWithNext=1,
    )
    code_font_size = 8.0
    # reportlab draws a paragraph's background OUTSIDE its measured box, so the vertical
    # padding eats into the surrounding space; keep it small and let spaceBefore/After
    # (which must exceed twice the padding) create the visible gap between blocks.
    code_pad_x, code_pad_y = 10.0, 7.0
    code_style = ParagraphStyle(
        "MdCode", parent=body_style, fontName=fonts["mono"], fontSize=code_font_size,
        leading=code_font_size * 1.45, textColor=colors.HexColor(CODE_FG),
        backColor=colors.HexColor(CODE_BG),
        borderPadding=(code_pad_y, code_pad_x, code_pad_y, code_pad_x),
        borderRadius=4, spaceBefore=15, spaceAfter=22, leftIndent=0, rightIndent=0,
    )
    quote_style = ParagraphStyle(
        "MdQuote", parent=body_style, leftIndent=14, textColor=colors.HexColor("#475569"),
        borderPadding=4, spaceBefore=4, spaceAfter=6,
    )
    caption_style = ParagraphStyle(
        "MdCaption", parent=body_style, fontSize=8, leading=10, alignment=1,
        textColor=colors.HexColor("#64748B"), spaceBefore=2, spaceAfter=8,
    )
    list_style = ParagraphStyle("MdListItem", parent=body_style, spaceAfter=2)

    pending = [dict(d) for d in (diagrams or [])]

    def take_diagram(code: str) -> dict | None:
        key = (code or "").strip()
        for i, d in enumerate(pending):
            if (d.get("code") or "").strip() == key:
                return pending.pop(i)
        return None

    def code_flowable(text: str, lang: str, style, chars: int):
        """``chars`` is the exact column count: DejaVu Sans Mono and Courier are 0.6em."""
        spans = _wrap_code_spans(_code_spans(text.replace("\t", "    "), lang), chars)
        return XPreformatted(_code_markup(spans) or " ", style)

    # A quote shifts its children right; nothing else about them changes, so the whole
    # renderer is re-entered with a larger indent rather than flattened into one string.
    quote_indent = 16.0

    def render(blocks: Sequence[Block], indent: float = 0.0, quoted: bool = False) -> list:
        avail = content_width - indent
        tag = f"{int(indent)}{'q' if quoted else ''}"
        b_style = ParagraphStyle(
            f"MdBody{tag}", parent=body_style, leftIndent=body_style.leftIndent + indent
        )
        h_style = ParagraphStyle(
            f"MdHead{tag}", parent=heading_style,
            leftIndent=heading_style.leftIndent + indent,
        )
        i_style = ParagraphStyle(f"MdItem{tag}", parent=list_style)
        c_style = ParagraphStyle(f"MdCodeI{tag}", parent=code_style, leftIndent=indent)
        cap_style = ParagraphStyle(
            f"MdCap{tag}", parent=caption_style, leftIndent=caption_style.leftIndent + indent
        )
        if quoted:
            b_style.textColor = quote_style.textColor
            i_style.textColor = quote_style.textColor
        chars = max(20, int((avail - 2 * code_pad_x) / (code_font_size * 0.602)))

        flow: list = []
        for block in blocks:
            kind = block[0]
            if kind == "h":
                size = max(15 - (block[1] - 1) * 1.6, 10.0)
                hs = ParagraphStyle(
                    f"MdH{block[1]}{tag}", parent=h_style, fontSize=size, leading=size + 4
                )
                flow.append(Paragraph(_inline_pdf(block[2]), hs))
            elif kind == "p":
                flow.append(Paragraph(_inline_pdf(block[1]), b_style))
            elif kind in ("ul", "ol"):
                items = [ListItem(Paragraph(_inline_pdf(it), i_style)) for it in block[1]]
                ordered = kind == "ol"
                flow.append(
                    ListFlowable(
                        items,
                        bulletType="1" if ordered else "bullet",
                        start=(block[2] if len(block) > 2 else 1) if ordered else None,
                        bulletFormat="%s." if ordered else None,
                        leftIndent=18 + indent,
                        bulletDedent=12,
                        bulletFontName=fonts["body"],
                        bulletFontSize=body_style.fontSize,
                        bulletColor=i_style.textColor,
                        spaceBefore=2,
                        spaceAfter=6,
                    )
                )
            elif kind == "code":
                text, lang = block[1], (block[2] if len(block) > 2 else "")
                if lang == "mermaid":
                    d = take_diagram(text)
                    card = (
                        _diagram_card(
                            d["data"], avail, 8.0 * 72,
                            natural_pt=float(d.get("width") or 0) * 0.75,
                        )
                        if d and d.get("data")
                        else None
                    )
                    if card is not None:
                        flow.append(card)
                        continue
                flow.append(code_flowable(text or " ", lang, c_style, chars))
            elif kind == "image":
                data = resolve_image_bytes(block[2])
                card = _diagram_card(data, avail, 7.0 * 72) if data else None
                if card is None:
                    flow.append(
                        Paragraph(_inline_pdf(f"[{block[1] or 'image'}]({block[2]})"), b_style)
                    )
                elif block[1]:
                    flow.append(
                        KeepTogether(
                            [card, Paragraph(escape(pdf_safe(block[1])), cap_style)]
                        )
                    )
                else:
                    flow.append(card)
            elif kind == "quote":
                flow.append(Spacer(1, 4))
                flow.extend(render(block[1], indent + quote_indent, quoted=True))
                flow.append(Spacer(1, 4))
            elif kind == "hr":
                flow.append(
                    HRFlowable(width=avail, thickness=0.6, color=colors.HexColor("#CBD5E1"),
                               spaceBefore=6, spaceAfter=6, hAlign="LEFT")
                )
            elif kind == "table":
                cols, rows = block[1], block[2]
                if not cols:
                    continue
                cell_style = ParagraphStyle(
                    "MdCell", parent=body_style, fontSize=8.5, leading=11, spaceAfter=0
                )
                head_style = ParagraphStyle(
                    "MdCellHead", parent=cell_style, fontName=fonts["bold"],
                    textColor=colors.HexColor("#1E1B4B"),
                )
                data = [[Paragraph(_inline_pdf(str(c)), head_style) for c in cols]]
                for r in rows:
                    data.append(
                        [
                            Paragraph(_inline_pdf(str(r[c_i])) if c_i < len(r) else "",
                                      cell_style)
                            for c_i in range(len(cols))
                        ]
                    )
                widths = _fit_col_widths(cols, rows, avail)
                # An indented table gets a blank spacer column so it lines up with the
                # quoted text around it; reportlab tables have no left indent of their own.
                x0 = 0
                if indent:
                    x0 = 1
                    widths = [indent] + widths
                    data = [[""] + row for row in data]
                # ``splitInRow`` lets a row taller than the page break across pages;
                # without it reportlab raises LayoutError and the export fails outright.
                tbl = Table(
                    data, colWidths=widths, hAlign="LEFT", repeatRows=1,
                    splitByRow=1, splitInRow=1,
                )

                tbl.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (x0, 0), (-1, 0), colors.HexColor("#EEF2FF")),
                            ("ROWBACKGROUNDS", (x0, 1), (-1, -1),
                             [colors.white, colors.HexColor("#F8FAFC")]),
                            ("GRID", (x0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 5),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                            ("TOPPADDING", (0, 0), (-1, -1), 4),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ]
                        + ([("LEFTPADDING", (0, 0), (0, -1), 0),
                            ("RIGHTPADDING", (0, 0), (0, -1), 0)] if x0 else [])
                    )
                )
                flow.append(tbl)
                flow.append(Spacer(1, 8))
        return flow

    return render(parse_blocks(md))
