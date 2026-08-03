from __future__ import annotations

import os
from xml.sax.saxutils import escape

from sqlalchemy import select
from sqlalchemy.orm import Session as DbSession

from .config import settings
from .models import Lane, LaneMessage, Provider, Session as ChatSession, Turn
from .tools.artifacts import generated_dir, new_stored_name, safe_download_name

_MIME = {
    "md": "text/markdown",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}

# How tall a prompt attachment may print. Wide enough to read, short enough that a photo
# doesn't push the answer onto the next page.
_ATTACHMENT_MAX_HEIGHT = 230.0


def _attachment_path(att) -> str:
    """Absolute path of an upload. ``storage_path`` is DB data, so it never leaves the dir."""
    return os.path.join(settings.UPLOAD_DIR, os.path.basename(att.storage_path))


def _turn_attachment_flowables(turn: Turn | None, content_width: float, caption_style) -> list:
    """Render what the user attached to a prompt, so the PDF shows the models' actual input.

    Images are embedded; documents are named (their text already reaches the models through
    the prompt itself). An image that can't be decoded degrades to its filename rather than
    failing the export.
    """
    from reportlab.platypus import Paragraph

    from .markdown_render import image_flowable, pdf_safe

    if turn is None or not turn.attachments:
        return []

    out: list = []
    named: list[str] = []
    for att in turn.attachments:
        card = None
        if att.kind == "image":
            try:
                with open(_attachment_path(att), "rb") as fh:
                    data = fh.read()
            except OSError:
                data = b""
            if data:
                card = image_flowable(data, content_width, _ATTACHMENT_MAX_HEIGHT)
        if card is None:
            named.append(att.filename)
            continue
        out.append(card)
        out.append(Paragraph(escape(pdf_safe(att.filename)), caption_style))

    if named:
        out.append(
            Paragraph(
                "Attached: " + escape(pdf_safe(", ".join(named))),
                caption_style,
            )
        )
    return out


def _add_turn_attachments_docx(doc, turn: Turn | None) -> None:
    """The Word counterpart of :func:`_turn_attachment_flowables`.

    python-docx only understands a handful of image formats, so anything it rejects (webp,
    typically) is converted to PNG first; whatever still won't go in is named instead.
    """
    from docx.shared import Emu, Inches, Pt, RGBColor

    if turn is None or not turn.attachments:
        return

    max_w, max_h = Inches(5.0), Inches(3.6)

    def caption(text: str) -> None:
        run = doc.add_paragraph().add_run(text)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    def place(path: str):
        try:
            return doc.add_picture(path)
        except Exception:  # noqa: BLE001 — unsupported format, try converting it
            import io

            from PIL import Image

            buf = io.BytesIO()
            with Image.open(path) as img:
                img.convert("RGB").save(buf, format="PNG")
            buf.seek(0)
            return doc.add_picture(buf)

    named: list[str] = []
    for att in turn.attachments:
        if att.kind != "image":
            named.append(att.filename)
            continue
        try:
            shape = place(_attachment_path(att))
        except Exception:  # noqa: BLE001 — missing or undecodable, name it instead
            named.append(att.filename)
            continue
        # add_picture uses the image's natural size, which can overflow the page.
        if shape.width > max_w:
            shape.height = Emu(int(shape.height * max_w / shape.width))
            shape.width = max_w
        if shape.height > max_h:
            shape.width = Emu(int(shape.width * max_h / shape.height))
            shape.height = max_h
        caption(att.filename)

    if named:
        caption("Attached: " + ", ".join(named))


def _gather(db: DbSession, session: ChatSession):
    lanes = [l for l in sorted(session.lanes, key=lambda x: x.position) if l.role == "responder"]
    turns = sorted(session.turns, key=lambda x: x.order_index)
    messages = db.scalars(
        select(LaneMessage)
        .join(Lane, Lane.id == LaneMessage.lane_id)
        .where(Lane.session_id == session.id, LaneMessage.role == "assistant")
    ).all()
    by_key: dict[tuple[str, str], LaneMessage] = {}
    for m in messages:
        by_key[(m.lane_id, m.turn_id)] = m
    providers = {p.id: p for p in db.scalars(select(Provider)).all()}

    def lane_label(lane: Lane) -> str:
        prov = providers.get(lane.provider_id)
        pname = prov.name if prov else "provider"
        return f"{lane.model} ({pname})"

    return lanes, turns, by_key, lane_label


def _export_markdown(db, session, path) -> None:
    lanes, turns, by_key, lane_label = _gather(db, session)
    out: list[str] = [f"# {session.title or 'Comparison'}", ""]
    for i, turn in enumerate(turns, 1):
        out.append(f"## Turn {i}")
        out.append("")
        out.append(f"**Prompt:** {turn.content}")
        out.append("")
        for lane in lanes:
            msg = by_key.get((lane.id, turn.id))
            if not msg or not msg.content:
                continue
            out.append(f"### {lane_label(lane)}")
            out.append("")
            out.append(msg.content)
            out.append("")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(out))


def _export_docx(db, session, path) -> None:
    from docx import Document
    from docx.shared import RGBColor

    from .markdown_render import render_markdown_docx

    lanes, turns, by_key, lane_label = _gather(db, session)
    doc = Document()
    h = doc.add_heading(session.title or "Comparison", level=0)
    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
    for i, turn in enumerate(turns, 1):
        doc.add_heading(f"Turn {i}", level=1)
        p = doc.add_paragraph()
        run = p.add_run("Prompt: ")
        run.bold = True
        p.add_run(turn.content or "")
        _add_turn_attachments_docx(doc, turn)
        for lane in lanes:
            msg = by_key.get((lane.id, turn.id))
            if not msg or not msg.content:
                continue
            hh = doc.add_heading(lane_label(lane), level=2)
            hh.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            render_markdown_docx(doc, msg.content, base_level=3)
    doc.save(path)


def _export_pdf(db, session, path) -> None:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    from .markdown_render import markdown_pdf_flowables

    lanes, turns, by_key, lane_label = _gather(db, session)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("T", parent=styles["Title"], textColor=HexColor("#1E1B4B"))
    turn_style = ParagraphStyle(
        "Turn", parent=styles["Heading1"], textColor=HexColor("#1E1B4B"), spaceBefore=12
    )
    lane_style = ParagraphStyle(
        "Lane", parent=styles["Heading3"], textColor=HexColor("#4F46E5"), spaceBefore=6
    )
    body = ParagraphStyle("Body", parent=styles["BodyText"], spaceAfter=5, leading=14)
    prompt_style = ParagraphStyle(
        "Prompt", parent=body, backColor=HexColor("#EEF2FF"), borderPadding=6, spaceAfter=8
    )
    caption_style = ParagraphStyle(
        "Caption", parent=body, fontSize=8, leading=11, alignment=1,
        textColor=HexColor("#64748B"), spaceBefore=2, spaceAfter=8,
    )

    doc = SimpleDocTemplate(
        path, pagesize=LETTER, topMargin=0.8 * inch, bottomMargin=0.8 * inch,
        leftMargin=0.8 * inch, rightMargin=0.8 * inch, title=session.title or "Comparison",
    )
    content_width = doc.width
    story: list = [Paragraph(escape(session.title or "Comparison"), title_style), Spacer(1, 8)]
    for i, turn in enumerate(turns, 1):
        story.append(Paragraph(f"Turn {i}", turn_style))
        story.append(Paragraph("<b>Prompt:</b> " + escape(turn.content or ""), prompt_style))
        story.extend(_turn_attachment_flowables(turn, content_width, caption_style))
        for lane in lanes:
            msg = by_key.get((lane.id, turn.id))
            if not msg or not msg.content:
                continue
            story.append(Paragraph(escape(lane_label(lane)), lane_style))
            story.extend(markdown_pdf_flowables(msg.content, body, content_width=content_width))
    doc.build(story)


_BUILDERS = {"md": _export_markdown, "docx": _export_docx, "pdf": _export_pdf}


def export_session(db: DbSession, session: ChatSession, fmt: str):
    """Export a whole session (all lanes side-by-side) to md/docx/pdf.

    Returns (stored_name, download_name, mime_type).
    """
    fmt = (fmt or "md").lower()
    builder = _BUILDERS.get(fmt)
    if not builder:
        raise ValueError(f"Unsupported export format: {fmt}")
    stored_name = new_stored_name(fmt)
    path = os.path.join(generated_dir(), stored_name)
    builder(db, session, path)
    download_name = safe_download_name(session.title or "comparison", fmt, fallback="comparison")
    return stored_name, download_name, _MIME.get(fmt, "application/octet-stream")


# ---------------------------------------------------------------------------
# Single answer export ("download this response as a PDF")
# ---------------------------------------------------------------------------

_PAGE_MARGIN = 54.0  # 0.75in — US Letter


def _numbered_canvas(footer_left: str, font: str):
    """Canvas subclass that stamps a rule + "Page X of Y" footer on every page.

    Page count is only known once the whole story is laid out, so pages are buffered and
    replayed on save.
    """
    from reportlab.lib.colors import HexColor
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas as pdf_canvas

    class NumberedCanvas(pdf_canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._pages: list[dict] = []

        def showPage(self):  # noqa: N802 — reportlab API
            self._pages.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            total = len(self._pages)
            for state in self._pages:
                self.__dict__.update(state)
                self._stamp(total)
                super().showPage()
            super().save()

        def _stamp(self, total: int) -> None:
            width, _ = LETTER
            self.saveState()
            self.setStrokeColor(HexColor("#E2E8F0"))
            self.setLineWidth(0.5)
            self.line(_PAGE_MARGIN, 44, width - _PAGE_MARGIN, 44)
            self.setFont(font, 7.5)
            self.setFillColor(HexColor("#94A3B8"))
            self.drawString(_PAGE_MARGIN, 32, footer_left[:110])
            self.drawRightString(
                width - _PAGE_MARGIN, 32, f"Page {self._pageNumber} of {total}"
            )
            self.restoreState()

    return NumberedCanvas


def _accent_card(text_flowables: list, width: float, bg: str, accent: str):
    """A tinted block with a coloured bar down its left edge."""
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    tbl = Table([["", text_flowables]], colWidths=[3.5, width - 3.5], hAlign="LEFT")
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor(accent)),
                ("BACKGROUND", (1, 0), (1, 0), colors.HexColor(bg)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (0, 0), 0),
                ("RIGHTPADDING", (0, 0), (0, 0), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("LEFTPADDING", (1, 0), (1, 0), 10),
                ("RIGHTPADDING", (1, 0), (1, 0), 10),
                ("TOPPADDING", (1, 0), (1, 0), 8),
                ("BOTTOMPADDING", (1, 0), (1, 0), 8),
            ]
        )
    )
    return tbl


def _answer_meta(message: LaneMessage) -> str:
    bits: list[str] = []
    if message.latency_ms is not None:
        bits.append(f"{message.latency_ms / 1000:.1f} s")
    if message.ttft_ms is not None:
        bits.append(f"{message.ttft_ms / 1000:.1f} s to first token")
    usage = message.usage_json or {}
    tokens = usage.get("completion_tokens")
    if tokens:
        bits.append(f"{tokens} tokens")
        if message.latency_ms:
            bits.append(f"{tokens / (message.latency_ms / 1000):.1f} tok/s")
    return "  \u00b7  ".join(bits)


def export_message_pdf(
    db: DbSession,
    session: ChatSession,
    message: LaneMessage,
    diagrams: list[dict] | None = None,
) -> tuple[str, str, str]:
    """Export ONE lane answer as a standalone US Letter PDF.

    ``diagrams`` carries images the chat UI already rendered for ```mermaid``` fences, so
    the document shows the same diagrams the user is looking at.

    Returns (stored_name, download_name, mime_type).
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from reportlab.platypus.flowables import HRFlowable

    from .markdown_render import markdown_pdf_flowables, pdf_fonts, pdf_safe

    fonts = pdf_fonts()
    lane = message.lane
    turn = message.turn
    provider = db.get(Provider, lane.provider_id) if lane and lane.provider_id else None
    model = (lane.model if lane else "") or "assistant"
    provider_name = provider.name if provider else "provider"
    title = session.title or "Chat response"

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "Base", parent=styles["BodyText"], fontName=fonts["body"], fontSize=10,
        leading=14.5, spaceAfter=6, textColor=colors.HexColor("#111827"),
    )
    brand_style = ParagraphStyle(
        "Brand", parent=base, fontName=fonts["bold"], fontSize=8, leading=10,
        textColor=colors.HexColor("#6366F1"), spaceAfter=2,
    )
    title_style = ParagraphStyle(
        "Title", parent=base, fontName=fonts["bold"], fontSize=19, leading=23,
        textColor=colors.HexColor("#1E1B4B"), spaceAfter=3,
    )
    meta_style = ParagraphStyle(
        "Meta", parent=base, fontSize=8.5, leading=12,
        textColor=colors.HexColor("#64748B"), spaceAfter=2,
    )
    label_style = ParagraphStyle(
        "Label", parent=base, fontName=fonts["bold"], fontSize=7.5, leading=10,
        textColor=colors.HexColor("#4F46E5"), spaceAfter=3,
    )
    prompt_style = ParagraphStyle(
        "PromptText", parent=base, fontSize=9.5, leading=13.5,
        textColor=colors.HexColor("#1F2937"), spaceAfter=0,
    )
    caption_style = ParagraphStyle(
        "Caption", parent=base, fontSize=8, leading=11, alignment=1,
        textColor=colors.HexColor("#64748B"), spaceBefore=2, spaceAfter=0,
    )

    stored_name = new_stored_name("pdf")
    path = os.path.join(generated_dir(), stored_name)
    doc = SimpleDocTemplate(
        path,
        pagesize=LETTER,
        topMargin=0.7 * 72,
        bottomMargin=0.85 * 72,
        leftMargin=_PAGE_MARGIN,
        rightMargin=_PAGE_MARGIN,
        title=f"{title} - {model}",
        author="MultiChat",
        subject=f"{model} ({provider_name})",
    )
    width = doc.width

    when = message.created_at.strftime("%d %b %Y, %H:%M UTC") if message.created_at else ""
    meta_bits = [b for b in (model, provider_name, when) if b]
    perf = _answer_meta(message)

    story: list = [
        Paragraph("MULTICHAT \u00b7 AI RESPONSE", brand_style),
        Paragraph(escape(pdf_safe(title)), title_style),
        Paragraph(escape(pdf_safe("  \u00b7  ".join(meta_bits))), meta_style),
    ]
    if perf:
        story.append(Paragraph(escape(pdf_safe(perf)), meta_style))
    story.append(Spacer(1, 12))

    # The prompt card carries whatever the user actually sent — the text AND any images or
    # documents attached to it, since those are as much a part of the question as the words.
    prompt_parts: list = []
    if turn is not None and (turn.content or "").strip():
        prompt_text = escape(pdf_safe(turn.content.strip())).replace("\n", "<br/>")
        prompt_parts.append(Paragraph(prompt_text, prompt_style))
    # 10pt of cell padding either side, plus the accent bar.
    prompt_parts.extend(_turn_attachment_flowables(turn, width - 30, caption_style))
    if prompt_parts:
        story.append(
            _accent_card(
                [Paragraph("PROMPT", label_style), *prompt_parts],
                width,
                bg="#EEF2FF",
                accent="#6366F1",
            )
        )
        story.append(Spacer(1, 14))

    story.append(Paragraph(f"RESPONSE \u2014 {escape(pdf_safe(model))}", label_style))
    story.append(
        HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#C7D2FE"),
                   spaceBefore=1, spaceAfter=8)
    )
    story.extend(
        markdown_pdf_flowables(
            message.content or "", base, diagrams=diagrams, content_width=width
        )
    )

    doc.build(story, canvasmaker=_numbered_canvas(f"{title} \u00b7 {model}", fonts["body"]))
    download_name = safe_download_name(f"{title}-{model}", "pdf", fallback="response")
    return stored_name, download_name, _MIME["pdf"]


# ---------------------------------------------------------------------------
# Deliberation export: the whole panel record, not just the answer
# ---------------------------------------------------------------------------


def export_deliberation_pdf(db: DbSession, run) -> tuple[str, str, str]:
    """Export a deliberation as a US Letter PDF: rounds, objections, synthesis, dissent.

    The point of the document is the middle section. Anyone can get an answer from one
    model; what this records is which claims were challenged, by whom, and what never got
    resolved.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import KeepTogether, Paragraph, SimpleDocTemplate, Spacer
    from reportlab.platypus.flowables import HRFlowable

    from .markdown_render import markdown_pdf_flowables, pdf_fonts, pdf_safe
    from .models import DeliberationStep, Session as ChatSession

    fonts = pdf_fonts()
    session = db.get(ChatSession, run.session_id)
    title = (session.title if session else None) or "Deliberation"
    lanes = {l.id: l for l in (session.lanes if session else [])}
    steps = db.scalars(
        select(DeliberationStep)
        .where(DeliberationStep.run_id == run.id)
        .order_by(DeliberationStep.round_index, DeliberationStep.created_at)
    ).all()

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "Base", parent=styles["BodyText"], fontName=fonts["body"], fontSize=10,
        leading=14.5, spaceAfter=6, textColor=colors.HexColor("#111827"),
    )
    brand = ParagraphStyle(
        "Brand", parent=base, fontName=fonts["bold"], fontSize=8, leading=10,
        textColor=colors.HexColor("#6366F1"), spaceAfter=2,
    )
    h_title = ParagraphStyle(
        "T", parent=base, fontName=fonts["bold"], fontSize=19, leading=23,
        textColor=colors.HexColor("#1E1B4B"), spaceAfter=3,
    )
    meta = ParagraphStyle(
        "M", parent=base, fontSize=8.5, leading=12,
        textColor=colors.HexColor("#64748B"), spaceAfter=2,
    )
    label = ParagraphStyle(
        "L", parent=base, fontName=fonts["bold"], fontSize=7.5, leading=10,
        textColor=colors.HexColor("#4F46E5"), spaceBefore=10, spaceAfter=3,
    )
    round_head = ParagraphStyle(
        "RH", parent=base, fontName=fonts["bold"], fontSize=12, leading=15,
        textColor=colors.HexColor("#1E1B4B"), spaceBefore=14, spaceAfter=4, keepWithNext=1,
    )
    model_head = ParagraphStyle(
        "MH", parent=base, fontName=fonts["bold"], fontSize=10, leading=13,
        textColor=colors.HexColor("#334155"), spaceBefore=8, spaceAfter=2,
    )
    reject = ParagraphStyle(
        "RJ", parent=base, fontSize=8.5, leading=11.5, leftIndent=10,
        textColor=colors.HexColor("#9F1239"), spaceAfter=1,
    )
    small = ParagraphStyle("S", parent=base, fontSize=8.5, leading=11.5, spaceAfter=2)

    stored_name = new_stored_name("pdf")
    path = os.path.join(generated_dir(), stored_name)
    doc = SimpleDocTemplate(
        path, pagesize=LETTER, topMargin=0.7 * 72, bottomMargin=0.85 * 72,
        leftMargin=_PAGE_MARGIN, rightMargin=_PAGE_MARGIN,
        title=f"Deliberation - {title}", author="MultiChat",
    )
    width = doc.width

    panel = [l for l in lanes.values() if l.role == "responder"]
    status_line = (
        f"{'converged' if run.converged else run.status.replace('_', ' ')} \u00b7 "
        f"{run.rounds_used} round(s) \u00b7 {run.total_calls} model calls \u00b7 "
        f"{run.wall_ms / 1000:.0f}s"
    )
    story: list = [
        Paragraph("MULTICHAT \u00b7 MODEL DELIBERATION", brand),
        Paragraph(escape(pdf_safe(title)), h_title),
        Paragraph(escape(pdf_safe(status_line)), meta),
        Paragraph(
            escape(pdf_safe("Panel: " + ", ".join(sorted(l.model for l in panel)))), meta
        ),
        Spacer(1, 12),
        _accent_card(
            [Paragraph("QUESTION", label), Paragraph(escape(pdf_safe(run.prompt or "")), base)],
            width, bg="#EEF2FF", accent="#6366F1",
        ),
        Spacer(1, 6),
    ]

    traces = {t.get("round"): t for t in (run.convergence_json or [])}
    rounds = sorted({s.round_index for s in steps if s.round_index < 90})
    last_round = rounds[-1] if rounds else 0
    for round_index in rounds:
        heading = (
            "Round 0 \u2014 independent drafts"
            if round_index == 0
            else f"Round {round_index} \u2014 peer review"
        )
        # Intermediate rounds carry only what changed. Reprinting every full answer for
        # every round triples the page count without adding audit value — the objections
        # are the record, and the final wording is shown in full below.
        full_body = round_index in (0, last_round)
        story.append(Paragraph(heading, round_head))
        if not full_body:
            story.append(
                Paragraph(
                    escape(pdf_safe("Objections raised this round; revised wording is shown "
                                    "in the final round below.")),
                    meta,
                )
            )
        for step in [s for s in steps if s.round_index == round_index]:
            if step.phase not in ("draft", "critique"):
                continue
            model = step.model or (lanes[step.lane_id].model if step.lane_id in lanes else "model")
            verdict = f"  [{step.verdict}]" if step.verdict else ""
            story.append(Paragraph(escape(pdf_safe(model + verdict)), model_head))
            if step.error:
                story.append(Paragraph(escape(pdf_safe(f"failed: {step.error}")), reject))
                continue
            output = step.output_json or {}
            rejected = [r for r in (output.get("rejected_claims") or []) if isinstance(r, dict)]
            if rejected:
                story.append(Paragraph("Rejected:", small))
                for item in rejected:
                    story.append(
                        Paragraph(
                            escape(pdf_safe(f"\u2717 {item.get('claim_id')} \u2014 {item.get('reason')}")),
                            reject,
                        )
                    )
            elif step.phase == "critique":
                story.append(Paragraph(escape(pdf_safe("No objections raised.")), small))
            body = str(output.get("revised_answer") or output.get("answer") or "")
            if body and full_body:
                story.extend(markdown_pdf_flowables(body, base, content_width=width))
        trace = traces.get(round_index)
        if trace:
            summary = (
                f"{len(trace.get('approvals') or [])}/{len(trace.get('responded') or [])} approved "
                f"\u00b7 {trace.get('open_objection_count', 0)} open objection(s) "
                f"\u00b7 claim overlap {round((trace.get('claim_overlap') or 0) * 100)}%"
            )
            story.append(
                _accent_card(
                    [Paragraph(escape(pdf_safe(summary)), small)],
                    width, bg="#F1F5F9", accent="#94A3B8",
                )
            )
            story.append(Spacer(1, 4))

    if run.synthesis:
        story.append(
            HRFlowable(width="100%", thickness=0.8, color=colors.HexColor("#C7D2FE"),
                       spaceBefore=14, spaceAfter=8)
        )
        story.append(Paragraph("SYNTHESIS", label))
        story.extend(markdown_pdf_flowables(run.synthesis, base, content_width=width))

    if run.minority_report:
        story.append(Spacer(1, 8))
        # The dissent is Markdown like any other model output — render it, don't dump it.
        story.append(
            _accent_card(
                [Paragraph("MINORITY REPORT \u2014 WHAT THE PANEL DID NOT SETTLE", label)]
                + markdown_pdf_flowables(
                    run.minority_report, small, content_width=width - 24
                ),
                width, bg="#FFFBEB", accent="#F59E0B",
            )
        )

    extraction = run.extraction_json or {}
    for heading, key in (("Do now", "do_now"), ("Consider later", "consider_later"), ("Skip", "skip")):
        items = [str(i) for i in (extraction.get(key) or []) if str(i).strip()]
        if not items:
            continue
        block = [Paragraph(heading.upper(), label)]
        block += [Paragraph(escape(pdf_safe(f"\u2022 {i}")), small) for i in items]
        story.append(KeepTogether(block))

    metrics = run.metrics_json or {}
    influence = metrics.get("influence") or {}
    capitulation = metrics.get("capitulation") or {}
    if influence or capitulation:
        rows = [Paragraph("PANEL METRICS", label)]
        for lane_id, value in sorted(influence.items(), key=lambda kv: -kv[1]):
            model = lanes[lane_id].model if lane_id in lanes else lane_id[:8]
            cap = capitulation.get(lane_id)
            text = f"{model} \u2014 influence {round(value * 100)}%"
            if cap is not None:
                text += f", capitulation {cap:.2f}"
            rows.append(Paragraph(escape(pdf_safe(text)), small))
        rows.append(
            Paragraph(
                escape(
                    pdf_safe(
                        "Influence = share of peer-accepted claims. Capitulation = changed "
                        "position without naming what changed its mind (lower is better)."
                    )
                ),
                meta,
            )
        )
        story.append(KeepTogether(rows))

    doc.build(story, canvasmaker=_numbered_canvas(f"Deliberation \u00b7 {title}", fonts["body"]))
    download_name = safe_download_name(f"deliberation-{title}", "pdf", fallback="deliberation")
    return stored_name, download_name, _MIME["pdf"]


# ---------------------------------------------------------------------------
# Deliberation exports in the other formats
# ---------------------------------------------------------------------------


def _deliberation_parts(db: DbSession, run):
    """Everything the non-PDF exporters need, in the order the panel produced it."""
    from .models import DeliberationStep

    session = db.get(ChatSession, run.session_id)
    title = (session.title if session else None) or "Deliberation"
    lanes = {l.id: l for l in (session.lanes if session else [])}
    steps = db.scalars(
        select(DeliberationStep)
        .where(DeliberationStep.run_id == run.id)
        .order_by(DeliberationStep.round_index, DeliberationStep.created_at)
    ).all()
    traces = {t.get("round"): t for t in (run.convergence_json or [])}
    return title, lanes, steps, traces


def _round_heading(round_index: int) -> str:
    return (
        "Round 0 — independent drafts"
        if round_index == 0
        else f"Round {round_index} — peer review"
    )


def _trace_summary(trace: dict) -> str:
    return (
        f"{len(trace.get('approvals') or [])}/{len(trace.get('responded') or [])} approved · "
        f"{trace.get('open_objection_count', 0)} open objection(s) · "
        f"claim overlap {round((trace.get('claim_overlap') or 0) * 100)}%"
    )


def _export_deliberation_markdown(db: DbSession, run, path: str) -> str:
    title, lanes, steps, traces = _deliberation_parts(db, run)
    panel = sorted(l.model for l in lanes.values() if l.role == "responder")
    out: list[str] = [
        f"# Deliberation — {title}",
        "",
        f"*{'converged' if run.converged else run.status.replace('_', ' ')} · "
        f"{run.rounds_used} round(s) · {run.total_calls} model calls · "
        f"{run.wall_ms / 1000:.0f}s*",
        "",
        f"**Panel:** {', '.join(panel)}",
        "",
        "## Question",
        "",
        run.prompt or "",
        "",
    ]
    rounds = sorted({s.round_index for s in steps if s.round_index < 90})
    last_round = rounds[-1] if rounds else 0
    for round_index in rounds:
        out += [f"## {_round_heading(round_index)}", ""]
        # Intermediate rounds carry only what changed — the objections are the record,
        # and the final wording is printed in full below.
        full_body = round_index in (0, last_round)
        for step in [s for s in steps if s.round_index == round_index]:
            if step.phase not in ("draft", "critique"):
                continue
            model = step.model or (lanes[step.lane_id].model if step.lane_id in lanes else "model")
            out.append(f"### {model}" + (f" — {step.verdict}" if step.verdict else ""))
            out.append("")
            if step.error:
                out += [f"> failed: {step.error}", ""]
                continue
            output = step.output_json or {}
            rejected = [r for r in (output.get("rejected_claims") or []) if isinstance(r, dict)]
            if rejected:
                out.append("**Rejected:**")
                out.append("")
                out += [
                    f"- ✗ {item.get('claim_id')} — {item.get('reason')}" for item in rejected
                ]
                out.append("")
            elif step.phase == "critique":
                out += ["No objections raised.", ""]
            body = str(output.get("revised_answer") or output.get("answer") or "")
            if body and full_body:
                out += [body, ""]
        trace = traces.get(round_index)
        if trace:
            out += [f"*{_trace_summary(trace)}*", ""]

    if run.synthesis:
        out += ["## Synthesis", "", run.synthesis, ""]
    if run.minority_report:
        out += [
            "## ⚠ Minority report — what the panel did not settle",
            "",
            run.minority_report,
            "",
        ]
    extraction = run.extraction_json or {}
    for heading, key in (("Do now", "do_now"), ("Consider later", "consider_later"), ("Skip", "skip")):
        items = [str(i) for i in (extraction.get(key) or []) if str(i).strip()]
        if items:
            out += [f"## {heading}", ""] + [f"- {i}" for i in items] + [""]

    metrics = run.metrics_json or {}
    influence = metrics.get("influence") or {}
    capitulation = metrics.get("capitulation") or {}
    if influence or capitulation:
        out += ["## Panel metrics", "", "| Model | Influence | Capitulation |", "| --- | --- | --- |"]
        for lane_id, value in sorted(influence.items(), key=lambda kv: -kv[1]):
            model = lanes[lane_id].model if lane_id in lanes else lane_id[:8]
            cap = capitulation.get(lane_id)
            out.append(f"| {model} | {round(value * 100)}% | {'—' if cap is None else f'{cap:.2f}'} |")
        out += [
            "",
            "*Influence = share of peer-accepted claims. Capitulation = changed position "
            "without naming what changed its mind (lower is better).*",
            "",
        ]
    with open(path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(out))
    return title


def _export_deliberation_docx(db: DbSession, run, path: str) -> str:
    from docx import Document
    from docx.shared import RGBColor

    from .markdown_render import render_markdown_docx

    title, lanes, steps, traces = _deliberation_parts(db, run)
    panel = sorted(l.model for l in lanes.values() if l.role == "responder")
    doc = Document()
    head = doc.add_heading(f"Deliberation — {title}", level=0)
    head.runs[0].font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)
    doc.add_paragraph(
        f"{'converged' if run.converged else run.status.replace('_', ' ')} · "
        f"{run.rounds_used} round(s) · {run.total_calls} model calls · "
        f"{run.wall_ms / 1000:.0f}s"
    )
    doc.add_paragraph("Panel: " + ", ".join(panel))
    doc.add_heading("Question", level=1)
    doc.add_paragraph(run.prompt or "")

    rounds = sorted({s.round_index for s in steps if s.round_index < 90})
    last_round = rounds[-1] if rounds else 0
    for round_index in rounds:
        doc.add_heading(_round_heading(round_index), level=1)
        full_body = round_index in (0, last_round)
        for step in [s for s in steps if s.round_index == round_index]:
            if step.phase not in ("draft", "critique"):
                continue
            model = step.model or (lanes[step.lane_id].model if step.lane_id in lanes else "model")
            sub = doc.add_heading(model + (f" — {step.verdict}" if step.verdict else ""), level=2)
            sub.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            if step.error:
                doc.add_paragraph(f"failed: {step.error}", style="Quote")
                continue
            output = step.output_json or {}
            rejected = [r for r in (output.get("rejected_claims") or []) if isinstance(r, dict)]
            for item in rejected:
                doc.add_paragraph(
                    f"✗ {item.get('claim_id')} — {item.get('reason')}", style="List Bullet"
                )
            if not rejected and step.phase == "critique":
                doc.add_paragraph("No objections raised.")
            body = str(output.get("revised_answer") or output.get("answer") or "")
            if body and full_body:
                render_markdown_docx(doc, body, base_level=3)
        trace = traces.get(round_index)
        if trace:
            doc.add_paragraph(_trace_summary(trace))

    if run.synthesis:
        doc.add_heading("Synthesis", level=1)
        render_markdown_docx(doc, run.synthesis, base_level=2)
    if run.minority_report:
        doc.add_heading("Minority report — what the panel did not settle", level=1)
        render_markdown_docx(doc, run.minority_report, base_level=2)
    extraction = run.extraction_json or {}
    for heading, key in (("Do now", "do_now"), ("Consider later", "consider_later"), ("Skip", "skip")):
        items = [str(i) for i in (extraction.get(key) or []) if str(i).strip()]
        if not items:
            continue
        doc.add_heading(heading, level=1)
        for item in items:
            doc.add_paragraph(item, style="List Bullet")
    doc.save(path)
    return title


def _export_deliberation_json(db: DbSession, run, path: str) -> str:
    """The full audit trail: every step's input, output, verdict and timing.

    This is the export that has no equivalent in a normal chat — it is what makes a
    surprising verdict reproducible.
    """
    import json

    title, lanes, steps, _ = _deliberation_parts(db, run)
    payload = {
        "id": run.id,
        "title": title,
        "status": run.status,
        "converged": run.converged,
        "prompt": run.prompt,
        "config": run.config_json or {},
        "rounds_used": run.rounds_used,
        "total_calls": run.total_calls,
        "wall_ms": run.wall_ms,
        "created_at": run.created_at.isoformat(),
        "panel": [
            {"lane_id": l.id, "model": l.model, "role": l.role}
            for l in sorted(lanes.values(), key=lambda x: x.position)
        ],
        "convergence": run.convergence_json or [],
        "vote": run.vote_json or {},
        "metrics": run.metrics_json or {},
        "synthesis": run.synthesis,
        "minority_report": run.minority_report,
        "extraction": run.extraction_json or {},
        "synthesis_critique": run.synthesis_critique_json or {},
        "steps": [
            {
                "id": s.id,
                "lane_id": s.lane_id,
                "message_id": s.message_id,
                "round": s.round_index,
                "phase": s.phase,
                "label": s.label,
                "model": s.model,
                "verdict": s.verdict,
                "input": s.input_json or {},
                "output": s.output_json or {},
                "degraded": s.degraded,
                "error": s.error,
                "latency_ms": s.latency_ms,
                "usage": s.usage_json,
                "created_at": s.created_at.isoformat(),
            }
            for s in steps
        ],
    }
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, indent=2, ensure_ascii=False)
    return title


_DELIBERATION_BUILDERS = {
    "md": _export_deliberation_markdown,
    "docx": _export_deliberation_docx,
    "json": _export_deliberation_json,
}

_DELIBERATION_MIME = {**_MIME, "json": "application/json"}


def export_deliberation(db: DbSession, run, fmt: str) -> tuple[str, str, str]:
    """Export a deliberation as pdf/md/docx/json.

    Returns (stored_name, download_name, mime_type).
    """
    fmt = (fmt or "pdf").lower()
    if fmt == "pdf":
        return export_deliberation_pdf(db, run)
    builder = _DELIBERATION_BUILDERS.get(fmt)
    if not builder:
        raise ValueError(f"Unsupported export format: {fmt}")
    stored_name = new_stored_name(fmt)
    path = os.path.join(generated_dir(), stored_name)
    title = builder(db, run, path)
    download_name = safe_download_name(f"deliberation-{title}", fmt, fallback="deliberation")
    return stored_name, download_name, _DELIBERATION_MIME.get(fmt, "application/octet-stream")
